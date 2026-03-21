#!/usr/bin/env python3
"""
Gmail Monitor for Railway — відстежує нові непрочитані листи,
витягує ВСІ 63-значні коди активації MS Office і надсилає в Telegram.
"""
import os, sys, re, json, base64, tempfile, time, warnings
warnings.filterwarnings("ignore")

import logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)]
)
log = logging.getLogger(__name__)

# ── Конфіг з env vars ──────────────────────────────────────────────────────────
TELEGRAM_TOKEN   = os.environ["TELEGRAM_TOKEN"]
TELEGRAM_CHAT_ID = os.environ["TELEGRAM_CHAT_ID"]
CHECK_INTERVAL   = int(os.environ.get("CHECK_INTERVAL", "30"))

# Credentials зберігаємо як файли (з env vars)
DATA_DIR         = os.environ.get("DATA_DIR", "/data")
CREDENTIALS_FILE = os.path.join(DATA_DIR, "credentials.json")
TOKEN_FILE       = os.path.join(DATA_DIR, "gmail_token.json")
PROCESSED_FILE   = os.path.join(DATA_DIR, "processed_ids.json")

os.makedirs(DATA_DIR, exist_ok=True)

# Записуємо credentials з env vars якщо файли не існують
if not os.path.exists(CREDENTIALS_FILE):
    creds_json = os.environ.get("GOOGLE_CREDENTIALS_JSON", "")
    if creds_json:
        with open(CREDENTIALS_FILE, "w") as f:
            f.write(creds_json)
        log.info("✅ credentials.json записано з env var")
    else:
        log.error("❌ GOOGLE_CREDENTIALS_JSON не встановлено!")
        sys.exit(1)

if not os.path.exists(TOKEN_FILE):
    token_json = os.environ.get("GOOGLE_TOKEN_JSON", "")
    if token_json:
        with open(TOKEN_FILE, "w") as f:
            f.write(token_json)
        log.info("✅ gmail_token.json записано з env var")
    else:
        log.error("❌ GOOGLE_TOKEN_JSON не встановлено!")
        sys.exit(1)

# ── Gmail API ──────────────────────────────────────────────────────────────────
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build

SCOPES = ["https://www.googleapis.com/auth/gmail.readonly"]

def get_gmail_service():
    creds = Credentials.from_authorized_user_file(TOKEN_FILE, SCOPES)
    if creds.expired and creds.refresh_token:
        creds.refresh(Request())
        with open(TOKEN_FILE, "w") as f:
            f.write(creds.to_json())
    return build("gmail", "v1", credentials=creds)


# ── Processed IDs ──────────────────────────────────────────────────────────────
def load_processed():
    if os.path.exists(PROCESSED_FILE):
        try:
            with open(PROCESSED_FILE) as f:
                return set(json.load(f))
        except Exception:
            pass
    return set()

def save_processed(ids: set):
    with open(PROCESSED_FILE, "w") as f:
        json.dump(list(ids), f)


# ── Пошук 63-значного коду ────────────────────────────────────────────────────
def normalize(text: str) -> str:
    text = text.replace("\n", " ")
    text = re.sub(r"[|/\\]", " ", text)
    text = re.sub(r"[^\d\s]", " ", text)
    return re.sub(r"\s+", " ", text).strip()

def find_activation_code(text: str):
    """Повертає перший знайдений код або None."""
    if not text:
        return None
    clean = normalize(text)
    m = re.search(r"\b(?:\d{7}\s+){8}\d{7}\b", clean)
    if m:
        return m.group(0).strip()
    m2 = re.search(r"(?<!\d)\d{63}(?!\d)", clean)
    if m2:
        d = m2.group(0)
        return " ".join(d[i:i+7] for i in range(0, 63, 7))
    return None


# ── OCR через Google Vision API (основний) + pytesseract (запасний) ───────────
VISION_API_KEY = os.environ.get("GOOGLE_VISION_API_KEY", "")

def ocr_via_google_vision(image_path: str) -> str:
    """OCR через Google Cloud Vision API — найточніший варіант."""
    if not VISION_API_KEY:
        return ""
    try:
        with open(image_path, "rb") as f:
            img_bytes = f.read()
        img_b64 = base64.b64encode(img_bytes).decode("utf-8")
        url = f"https://vision.googleapis.com/v1/images:annotate?key={VISION_API_KEY}"
        payload = {"requests": [{"image": {"content": img_b64},
                                 "features": [{"type": "TEXT_DETECTION"}]}]}
        resp = req.post(url, json=payload, timeout=30)
        resp.raise_for_status()
        annotations = resp.json().get("responses", [{}])[0].get("textAnnotations", [])
        if annotations:
            text = annotations[0].get("description", "")
            log.info("  OCR: Google Vision API успішно")
            return text
    except Exception as e:
        log.warning(f"Google Vision API помилка: {e}")
    return ""

def _preprocess(img, strategy: int):
    """Різні стратегії обробки зображення для pytesseract."""
    from PIL import ImageOps, ImageFilter, ImageEnhance
    img = img.convert("L")
    if strategy == 0:
        img = ImageOps.autocontrast(img)
        img = img.filter(ImageFilter.SHARPEN)
        w, h = img.size
        img = img.resize((w * 2, h * 2))
        img = img.point(lambda p: 255 if p > 160 else 0)
    elif strategy == 1:
        img = ImageOps.autocontrast(img, cutoff=2)
        img = ImageEnhance.Contrast(img).enhance(3.0)
        img = ImageEnhance.Sharpness(img).enhance(3.0)
        w, h = img.size
        img = img.resize((w * 3, h * 3))
        import statistics
        threshold = statistics.median(list(img.getdata()))
        img = img.point(lambda p: 255 if p > threshold else 0)
    elif strategy == 2:
        img = ImageOps.invert(img)
        img = ImageOps.autocontrast(img, cutoff=5)
        img = ImageEnhance.Contrast(img).enhance(2.0)
        w, h = img.size
        img = img.resize((w * 3, h * 3))
        img = img.point(lambda p: 255 if p > 128 else 0)
    elif strategy == 3:
        img = ImageOps.autocontrast(img, cutoff=1)
        img = ImageEnhance.Sharpness(img).enhance(2.0)
        w, h = img.size
        img = img.resize((w * 3, h * 3))
    return img

def ocr_via_tesseract(image_path: str) -> str:
    """Запасний OCR через pytesseract з кількома стратегіями обробки."""
    try:
        import pytesseract
        from PIL import Image
        base_img = Image.open(image_path)
        cfg = "--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789 "
        best_text = ""
        for strategy in range(4):
            try:
                processed = _preprocess(base_img.copy(), strategy)
                text = pytesseract.image_to_string(processed, config=cfg)
                if find_activation_code(text):
                    log.info(f"  OCR: tesseract стратегія {strategy} знайшла код")
                    return text
                if sum(c.isdigit() for c in text) > sum(c.isdigit() for c in best_text):
                    best_text = text
            except Exception:
                continue
        return best_text
    except Exception as e:
        log.warning(f"Tesseract OCR помилка: {e}")
        return ""

def ocr_image(image_path: str) -> str:
    """Спочатку Google Vision API, при невдачі — pytesseract."""
    # Спробуємо Google Vision API
    text = ocr_via_google_vision(image_path)
    if text:
        return text
    # Запасний варіант — pytesseract
    log.info("  OCR: fallback на pytesseract")
    return ocr_via_tesseract(image_path)


# ── Декодування тіла листа ────────────────────────────────────────────────────
def decode_b64(data: str) -> str:
    try:
        missing = len(data) % 4
        if missing:
            data += "=" * (4 - missing)
        return base64.urlsafe_b64decode(data).decode("utf-8", errors="ignore")
    except Exception:
        return ""

def get_body_text(payload: dict) -> str:
    mime = payload.get("mimeType", "")
    data = payload.get("body", {}).get("data", "")
    if mime in ("text/plain", "text/html") and data:
        return decode_b64(data)
    return "".join(get_body_text(p) for p in payload.get("parts", []))

def get_plain_body(payload: dict) -> str:
    """Повертає чистий текст листа (без HTML тегів і зайвих пробілів)."""
    # Спочатку шукаємо text/plain
    def _find(parts, target_mime):
        for part in parts:
            if part.get("mimeType") == target_mime:
                data = part.get("body", {}).get("data", "")
                if data:
                    return decode_b64(data)
            sub = part.get("parts", [])
            if sub:
                result = _find(sub, target_mime)
                if result:
                    return result
        return ""

    text = _find(payload.get("parts", [payload]), "text/plain")
    if not text:
        html = _find(payload.get("parts", [payload]), "text/html")
        if html:
            # Прибираємо HTML теги
            text = re.sub(r"<[^>]+>", " ", html)

    # Чистимо текст
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = text.strip()

    # Обрізаємо якщо дуже довгий (Telegram обмеження 4096 символів)
    if len(text) > 600:
        text = text[:600].rsplit("\n", 1)[0] + "\n..."

    return text


# ── Обробка вкладень — повертає СПИСОК всіх знайдених кодів ──────────────────
def get_attachment_bytes(service, msg_id, part):
    body = part.get("body", {})
    inline = body.get("data")
    if inline:
        missing = len(inline) % 4
        return base64.urlsafe_b64decode(inline + "=" * missing)
    att_id = body.get("attachmentId")
    if att_id:
        att = service.users().messages().attachments().get(
            userId="me", messageId=msg_id, id=att_id).execute()
        data = att.get("data", "")
        missing = len(data) % 4
        return base64.urlsafe_b64decode(data + "=" * missing)
    return None

def extract_codes_from_docx(docx_bytes: bytes) -> list:
    """Витягує всі коди з Word (.docx) документа."""
    codes = []
    try:
        import docx, io
        doc = docx.Document(io.BytesIO(docx_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)
        # також перевіряємо таблиці
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        code = find_activation_code(text)
        if code:
            codes.append(code)
    except Exception as e:
        log.warning(f"DOCX помилка: {e}")
    return codes

def extract_codes_from_xlsx(xlsx_bytes: bytes) -> list:
    """Витягує всі коди з Excel (.xlsx) файлу."""
    codes = []
    try:
        import openpyxl, io
        wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), read_only=True, data_only=True)
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                text = " ".join(str(c) for c in row if c is not None)
                code = find_activation_code(text)
                if code and code not in codes:
                    codes.append(code)
    except Exception as e:
        log.warning(f"XLSX помилка: {e}")
    return codes

def extract_codes_from_pdf(pdf_bytes: bytes) -> list:
    """Витягує ВСІ коди з PDF (текст + зображення)."""
    codes = []
    try:
        import fitz
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page in doc:
            code = find_activation_code(page.get_text())
            if code and code not in codes:
                codes.append(code)
            for img_info in page.get_images(full=True):
                base_img = doc.extract_image(img_info[0])
                with tempfile.NamedTemporaryFile(suffix=f".{base_img['ext']}", delete=False) as tmp:
                    tmp.write(base_img["image"])
                    tmp_path = tmp.name
                try:
                    code = find_activation_code(ocr_image(tmp_path))
                    if code and code not in codes:
                        codes.append(code)
                finally:
                    os.unlink(tmp_path)
    except Exception as e:
        log.warning(f"PDF помилка: {e}")
    return codes

def process_attachments(service, msg_id: str, payload: dict) -> list:
    """Обходить всі вкладення і повертає СПИСОК всіх знайдених кодів."""
    codes = []

    def _walk(parts):
        for part in parts:
            mime     = part.get("mimeType", "")
            filename = (part.get("filename") or "").lower()
            is_image = (
                mime.startswith("image/") or
                any(filename.endswith(ext) for ext in (".png",".jpg",".jpeg",".bmp",".tiff",".webp"))
            )
            if is_image:
                img_bytes = get_attachment_bytes(service, msg_id, part)
                if img_bytes:
                    ext = mime.split("/")[-1].split(";")[0] if "/" in mime else "png"
                    with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
                        tmp.write(img_bytes)
                        tmp_path = tmp.name
                    try:
                        code = find_activation_code(ocr_image(tmp_path))
                        if code and code not in codes:
                            codes.append(code)
                    finally:
                        os.unlink(tmp_path)
            elif mime == "application/pdf" or filename.endswith(".pdf"):
                pdf_bytes = get_attachment_bytes(service, msg_id, part)
                if pdf_bytes:
                    for code in extract_codes_from_pdf(pdf_bytes):
                        if code not in codes:
                            codes.append(code)
            elif (mime in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           "application/msword") or
                  filename.endswith((".docx", ".doc"))):
                docx_bytes = get_attachment_bytes(service, msg_id, part)
                if docx_bytes:
                    for code in extract_codes_from_docx(docx_bytes):
                        if code not in codes:
                            codes.append(code)
            elif (mime in ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                           "application/vnd.ms-excel") or
                  filename.endswith((".xlsx", ".xls"))):
                xlsx_bytes = get_attachment_bytes(service, msg_id, part)
                if xlsx_bytes:
                    for code in extract_codes_from_xlsx(xlsx_bytes):
                        if code not in codes:
                            codes.append(code)
            sub = part.get("parts", [])
            if sub:
                _walk(sub)

    _walk(payload.get("parts", []))
    return codes


# ── Telegram ───────────────────────────────────────────────────────────────────
import requests as req

def send_telegram(text: str):
    url = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendMessage"
    req.post(url, data={"chat_id": TELEGRAM_CHAT_ID, "text": text, "parse_mode": "Markdown"}, timeout=20)

def notify(sender_email: str, subject: str, codes: list, msg_id: str = "", body_text: str = ""):
    """Надсилає всі знайдені коди — кожен окремим повідомленням."""
    gmail_url = f"https://mail.google.com/mail/u/0/#inbox/{msg_id}" if msg_id else ""
    email_link = f"[{sender_email}]({gmail_url})" if gmail_url else f"`{sender_email}`"
    body_section = f"\n💬 *Повідомлення:*\n{body_text}\n" if body_text else ""
    for i, code in enumerate(codes, 1):
        num = f" #{i}" if len(codes) > 1 else ""
        msg = (
            f"📧 *Від:* {email_link}\n"
            f"📌 *Тема:* {subject or '—'}\n"
            f"{body_section}\n"
            f"🔑 *Код активації{num}:*\n`{code}`"
        )
        send_telegram(msg)
        log.info(f"✅ Відправлено код{num}: {code}")


# ── Одна перевірка ────────────────────────────────────────────────────────────
def check_once(service, processed: set) -> set:
    result = service.users().messages().list(
        userId="me", q="is:unread", maxResults=50).execute()
    messages = result.get("messages", [])
    if not messages:
        return processed

    new_messages = [m for m in reversed(messages) if m["id"] not in processed]
    if not new_messages:
        return processed

    log.info(f"📬 Нових листів: {len(new_messages)}")

    for msg_ref in new_messages:
        msg_id = msg_ref["id"]
        msg = service.users().messages().get(
            userId="me", id=msg_id, format="full").execute()

        headers      = {h["name"]: h["value"] for h in msg["payload"]["headers"]}
        sender       = headers.get("From", "невідомо")
        subject      = headers.get("Subject", "")
        m            = re.search(r'<([^>]+)>', sender)
        sender_email = m.group(1) if m else sender

        log.info(f"→ {sender_email} | {subject}")

        # Отримуємо текст листа для відображення
        body_text = get_plain_body(msg["payload"])

        # Збираємо всі коди: спочатку з тіла, потім з вкладень
        codes = []
        body_code = find_activation_code(get_body_text(msg["payload"]))
        if body_code:
            codes.append(body_code)

        for code in process_attachments(service, msg_id, msg["payload"]):
            if code not in codes:
                codes.append(code)

        processed.add(msg_id)

        if codes:
            log.info(f"  ↳ Знайдено кодів: {len(codes)}")
            notify(sender_email, subject, codes, msg_id, body_text)
        else:
            log.info(f"  ↳ Код не знайдено — ігнорується")

    save_processed(processed)
    return processed


# ── Головний цикл ─────────────────────────────────────────────────────────────
def main():
    log.info("🚀 Gmail Monitor запущено на Railway")
    service   = get_gmail_service()
    processed = load_processed()

    # Перший запуск — зберігаємо всі поточні unread IDs без обробки
    if not os.path.exists(PROCESSED_FILE):
        result = service.users().messages().list(
            userId="me", q="is:unread", maxResults=500).execute()
        existing = {m["id"] for m in result.get("messages", [])}
        save_processed(existing)
        processed = existing
        log.info(f"Перший запуск. Збережено {len(existing)} існуючих листів.")
        send_telegram("🚀 Gmail моніторинг запущено! Чекаю нових листів з кодами активації.")

    last_error_notify = 0  # час останнього Telegram-повідомлення про помилку
    ERROR_NOTIFY_INTERVAL = 1800  # не спамити частіше ніж раз на 30 хвилин

    while True:
        try:
            processed = check_once(service, processed)
            last_error_notify = 0  # скидаємо лічильник якщо все ок
        except Exception as e:
            log.error(f"Помилка: {e}")

            # Надсилаємо в Telegram не частіше ніж раз на 30 хвилин
            now = time.time()
            if now - last_error_notify > ERROR_NOTIFY_INTERVAL:
                try:
                    send_telegram(f"⚠️ *Gmail Monitor: помилка!*\n`{str(e)[:300]}`\n\nСпробую перепідключитись...")
                    last_error_notify = now
                except Exception:
                    pass

            # Спроба перепідключитись до Gmail
            try:
                service = get_gmail_service()
                log.info("✅ Перепідключення до Gmail успішне")
            except Exception as e2:
                log.error(f"Перепідключення не вдалось: {e2}")

        time.sleep(CHECK_INTERVAL)

if __name__ == "__main__":
    main()
